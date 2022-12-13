# pip install python-docx
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE



if __name__ == '__main__':
    path = "pois.docx"
    doc = docx.Document(path)
    text = []
    i = 0

    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(12)  #Кегль
    font_object.name = 'Times New Roman' #Шрифт

    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

        if "Рисунок" in paragraph.text: # если в строке есть слово Рисунок
            i += 1
            s = str(paragraph.text)
            index = 0 #индекс второго пробела в строке
            k = 0
            for j in range(len(s)):
                if s[j] == " ":
                    k += 1
                    if k == 2:
                        index = j
                        break
            s1 = "Рисунок " + str(i) + ' ' + s[index:] #меняем строку
            paragraph.clear()
            paragraph.add_run(s1, style='CommentsStyle').bold = True #вставляем текс, выделяем жирным


    doc.save(path[:6] + "_fixed.docx")


